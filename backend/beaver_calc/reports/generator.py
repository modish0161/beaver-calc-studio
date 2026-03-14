"""
Report Generation Module for BeaverCalc Studio
Generates PDF and DOCX reports with Beaver Bridges branding
"""

import os
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional
import io

from jinja2 import Environment, FileSystemLoader
try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except (OSError, ImportError):
    HTML = None
    CSS = None
    WEASYPRINT_AVAILABLE = False
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Beaver Bridges brand colors
BRAND_PRIMARY = RGBColor(38, 44, 83)  # #262C53
BRAND_TURQUOISE = RGBColor(0, 217, 255)  # #00D9FF
BRAND_SILVER = RGBColor(192, 192, 192)  # #C0C0C0

TEMPLATE_DIR = Path(__file__).parent / 'templates'


class ReportGenerator:
    """Generate PDF and DOCX calculation reports"""
    
    def __init__(self):
        self.env = Environment(loader=FileSystemLoader(str(TEMPLATE_DIR)))
        
    def generate_pdf(
        self,
        run_data: Dict[str, Any],
        output_path: Optional[str] = None
    ) -> bytes:
        """
        Generate PDF report from calculation run data
        
        Args:
            run_data: Dictionary containing calculation data
            output_path: Optional file path to save PDF
            
        Returns:
            PDF content as bytes
        """
        if not WEASYPRINT_AVAILABLE:
            raise RuntimeError("PDF generation requires weasyprint with GTK libraries installed")
        template = self.env.get_template('calculation_report.html')
        
        # Prepare context data
        context = self._prepare_context(run_data)
        
        # Render HTML
        html_content = template.render(**context)
        
        # Generate PDF
        pdf_bytes = HTML(string=html_content).write_pdf()
        
        # Save to file if path provided
        if output_path:
            with open(output_path, 'wb') as f:
                f.write(pdf_bytes)
                
        return pdf_bytes
    
    def generate_docx(
        self,
        run_data: Dict[str, Any],
        output_path: Optional[str] = None
    ) -> bytes:
        """
        Generate DOCX report from calculation run data
        
        Args:
            run_data: Dictionary containing calculation data
            output_path: Optional file path to save DOCX
            
        Returns:
            DOCX content as bytes
        """
        doc = Document()
        context = self._prepare_context(run_data)
        
        # Configure document styles
        self._setup_styles(doc)
        
        # Add header
        self._add_header(doc, context)
        
        # Add title section
        self._add_title_section(doc, context)
        
        # Add project information
        self._add_project_info(doc, context)
        
        # Add scope & assumptions
        self._add_scope_and_assumptions(doc, context)
        
        # Add input parameters
        self._add_input_parameters(doc, context)
        
        # Add calculation steps
        self._add_calculation_steps(doc, context)
        
        # Add results
        self._add_results(doc, context)
        
        # Add design checks
        self._add_design_checks(doc, context)
        
        # Add sensitivity & warnings
        self._add_sensitivity_warnings(doc, context)
        
        # Add design summary
        self._add_design_summary(doc, context)
        
        # Add conclusions & sign-off block
        self._add_conclusions_signoff(doc, context)
        
        # Add disclaimer
        self._add_disclaimer(doc, context)
        
        # Add appendices
        self._add_appendices(doc, context)
        
        # Add footer
        self._add_footer(doc, context)
        
        # Save to bytes
        docx_io = io.BytesIO()
        doc.save(docx_io)
        docx_bytes = docx_io.getvalue()
        
        # Save to file if path provided
        if output_path:
            with open(output_path, 'wb') as f:
                f.write(docx_bytes)
                
        return docx_bytes
    
    def _prepare_context(self, run_data: Dict[str, Any]) -> Dict[str, Any]:
        """Prepare template context from run data"""
        current_date = datetime.now().strftime('%d %B %Y')
        current_year = datetime.now().year
        
        context = {
            'current_date': current_date,
            'current_year': current_year,
            'run_id': run_data.get('id', 'N/A'),
            'version': '1.0.0',
            
            # Calculator info
            'calculator_name': run_data.get('calculator_name', 'Structural Calculation'),
            'calculator_description': run_data.get('calculator_description', ''),
            
            # Project info
            'project_name': run_data.get('project_name', 'Unnamed Project'),
            'project_id': run_data.get('project_id', 'N/A'),
            'client_name': run_data.get('client_name', 'N/A'),
            'engineer_name': run_data.get('engineer_name', 'BeaverCalc User'),
            'design_code': run_data.get('design_code', 'EN 1993-1-1'),
            'calc_date': run_data.get('created_at', current_date),
            
            # Calculation data
            'input_parameters': run_data.get('inputs', []),
            'calculation_steps': run_data.get('steps', []),
            'results': run_data.get('results', []),
            'design_checks': run_data.get('checks', []),
            'design_summary': run_data.get('summary', 'Calculation completed successfully.'),
        }
        
        return context
    
    def _setup_styles(self, doc: Document):
        """Configure document styles with Beaver Bridges branding"""
        styles = doc.styles
        
        # Heading 1 - Main section headers
        if 'Heading 1' not in styles:
            heading1 = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        else:
            heading1 = styles['Heading 1']
        heading1.font.name = 'Arial'
        heading1.font.size = Pt(16)
        heading1.font.bold = True
        heading1.font.color.rgb = BRAND_PRIMARY
        
        # Heading 2 - Subsection headers
        if 'Heading 2' not in styles:
            heading2 = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        else:
            heading2 = styles['Heading 2']
        heading2.font.name = 'Arial'
        heading2.font.size = Pt(13)
        heading2.font.bold = True
        heading2.font.color.rgb = BRAND_TURQUOISE
    
    def _add_header(self, doc: Document, context: Dict[str, Any]):
        """Add report header with logo and info"""
        # Logo section
        header = doc.add_paragraph()
        header_run = header.add_run('BEAVER BRIDGES')
        header_run.font.size = Pt(24)
        header_run.font.bold = True
        header_run.font.color.rgb = BRAND_PRIMARY
        
        tagline = doc.add_paragraph('Structural Engineering Excellence')
        tagline_run = tagline.runs[0]
        tagline_run.font.size = Pt(9)
        tagline_run.font.color.rgb = BRAND_TURQUOISE
        tagline_run.font.bold = True
        
        # Report info
        info = doc.add_paragraph()
        info.add_run(f"Report Generated: {context['current_date']}\n")
        info.add_run(f"Report ID: {context['run_id']}\n")
        info.add_run(f"Version: {context['version']}")
        info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in info.runs:
            run.font.size = Pt(8)
        
        doc.add_paragraph()  # Spacing
    
    def _add_title_section(self, doc: Document, context: Dict[str, Any]):
        """Add calculation title section"""
        title = doc.add_heading(context['calculator_name'], level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph(context['calculator_description'])
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.color.rgb = BRAND_TURQUOISE
        subtitle_run.font.size = Pt(11)
        
        doc.add_paragraph()  # Spacing
    
    def _add_project_info(self, doc: Document, context: Dict[str, Any]):
        """Add project information section"""
        doc.add_heading('PROJECT INFORMATION', level=2)
        
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'
        
        info_items = [
            ('Project Name', context['project_name']),
            ('Project ID', context['project_id']),
            ('Client', context['client_name']),
            ('Engineer', context['engineer_name']),
            ('Design Code', context['design_code']),
            ('Calculation Date', context['calc_date']),
        ]
        
        for i, (label, value) in enumerate(info_items):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)
            table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()  # Spacing
    
    def _add_input_parameters(self, doc: Document, context: Dict[str, Any]):
        """Add input parameters table"""
        doc.add_heading('INPUT PARAMETERS', level=2)
        
        params = context.get('input_parameters', [])
        if not params:
            doc.add_paragraph('No input parameters provided.')
            return
        
        table = doc.add_table(rows=len(params) + 1, cols=5)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        headers = ['Parameter', 'Symbol', 'Value', 'Unit', 'Reference']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Data rows
        for i, param in enumerate(params, 1):
            table.rows[i].cells[0].text = param.get('name', '')
            table.rows[i].cells[1].text = param.get('symbol', '')
            table.rows[i].cells[2].text = str(param.get('value', ''))
            table.rows[i].cells[3].text = param.get('unit', '')
            table.rows[i].cells[4].text = param.get('reference', '')
        
        doc.add_paragraph()  # Spacing
    
    def _add_calculation_steps(self, doc: Document, context: Dict[str, Any]):
        """Add calculation methodology section"""
        doc.add_heading('CALCULATION METHODOLOGY', level=2)
        
        steps = context.get('calculation_steps', [])
        if not steps:
            doc.add_paragraph('No calculation steps provided.')
            return
        
        for i, step in enumerate(steps, 1):
            step_para = doc.add_paragraph()
            step_para.add_run(f"Step {i}: {step.get('title', '')}\n").font.bold = True
            step_para.add_run(step.get('description', ''))
            
            if 'equation' in step:
                eq_para = doc.add_paragraph(step['equation'])
                eq_para.runs[0].font.name = 'Courier New'
                eq_para.runs[0].font.size = Pt(9)
            
            if 'code_reference' in step:
                ref = step['code_reference']
                ref_para = doc.add_paragraph()
                ref_para.add_run(f"{ref.get('clause', '')}: ").font.bold = True
                ref_para.add_run(ref.get('description', ''))
                ref_para.runs[0].font.color.rgb = BRAND_TURQUOISE
        
        doc.add_paragraph()  # Spacing
    
    def _add_results(self, doc: Document, context: Dict[str, Any]):
        """Add calculation results section"""
        doc.add_heading('CALCULATION RESULTS', level=2)
        
        results = context.get('results', [])
        if not results:
            doc.add_paragraph('No results available.')
            return
        
        for result in results:
            result_para = doc.add_paragraph()
            result_para.add_run(f"{result.get('name', '')}: ").font.bold = True
            value_run = result_para.add_run(f"{result.get('value', '')} {result.get('unit', '')}")
            value_run.font.color.rgb = BRAND_TURQUOISE
            value_run.font.size = Pt(12)
            value_run.font.bold = True
        
        doc.add_paragraph()  # Spacing
    
    def _add_design_checks(self, doc: Document, context: Dict[str, Any]):
        """Add design checks table"""
        checks = context.get('design_checks', [])
        if not checks:
            return
        
        doc.add_heading('DESIGN CHECKS', level=2)
        
        table = doc.add_table(rows=len(checks) + 1, cols=5)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        headers = ['Check', 'Capacity', 'Demand', 'Utilization', 'Status']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
        
        # Data rows
        for i, check in enumerate(checks, 1):
            table.rows[i].cells[0].text = check.get('name', '')
            table.rows[i].cells[1].text = f"{check.get('capacity', '')} {check.get('unit', '')}"
            table.rows[i].cells[2].text = f"{check.get('demand', '')} {check.get('unit', '')}"
            table.rows[i].cells[3].text = f"{check.get('utilization', '')}%"
            
            status = check.get('status', '')
            status_cell = table.rows[i].cells[4]
            status_cell.text = status
            if status == 'PASS':
                status_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(40, 167, 69)
            elif status == 'FAIL':
                status_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(220, 53, 69)
        
        doc.add_paragraph()  # Spacing
    
    def _add_design_summary(self, doc: Document, context: Dict[str, Any]):
        """Add design summary section"""
        doc.add_heading('DESIGN SUMMARY', level=2)
        doc.add_paragraph(context.get('design_summary', ''))
        doc.add_paragraph()  # Spacing
    
    def _add_scope_and_assumptions(self, doc: Document, context: Dict[str, Any]):
        """Add scope & assumptions section (Appendix B requirement)"""
        doc.add_heading('SCOPE & ASSUMPTIONS', level=2)

        scope_items = context.get('scope_assumptions', [
            'This calculation covers the structural adequacy of the element described above.',
            'All loads and load combinations are as specified in the input parameters.',
            'Material properties are in accordance with the referenced design code.',
            'The structure is assumed to be adequately braced out-of-plane unless otherwise noted.',
            'Connections and fixings are outside the scope of this calculation unless explicitly checked.',
            'Site conditions and geotechnical parameters are as provided by the client/geotechnical engineer.',
        ])

        for item in scope_items:
            p = doc.add_paragraph(item, style='List Bullet')
            for run in p.runs:
                run.font.size = Pt(10)

        doc.add_paragraph()  # Spacing

    def _add_sensitivity_warnings(self, doc: Document, context: Dict[str, Any]):
        """Add sensitivity analysis and warnings section"""
        doc.add_heading('SENSITIVITY & WARNINGS', level=2)

        # Check for high utilisations
        checks = context.get('design_checks', [])
        warnings = []
        for check in checks:
            util = check.get('utilization', 0)
            if isinstance(util, (int, float)):
                if util > 90:
                    warnings.append(
                        f"High utilisation ({util:.0f}%) on '{check.get('name', 'check')}' — "
                        f"consider increasing member size or reducing load."
                    )
                if check.get('status') == 'FAIL':
                    warnings.append(
                        f"FAIL on '{check.get('name', 'check')}' — design does not satisfy code requirements."
                    )

        if not warnings:
            warnings = ['No critical sensitivity issues identified. All checks within acceptable limits.']

        for w in warnings:
            p = doc.add_paragraph(w, style='List Bullet')
            for run in p.runs:
                run.font.size = Pt(10)

        doc.add_paragraph()  # Spacing

    def _add_conclusions_signoff(self, doc: Document, context: Dict[str, Any]):
        """Add conclusions and sign-off block"""
        doc.add_heading('CONCLUSIONS & SIGN-OFF', level=2)

        # Conclusions
        summary = context.get('design_summary', 'Calculation completed.')
        doc.add_paragraph(summary)
        doc.add_paragraph()

        # Sign-off table
        table = doc.add_table(rows=4, cols=4)
        table.style = 'Light Grid Accent 1'

        headers = ['Role', 'Name', 'Signature', 'Date']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True

        roles = ['Designed by', 'Checked by', 'Approved by']
        for i, role in enumerate(roles, 1):
            table.rows[i].cells[0].text = role
            if role == 'Designed by':
                table.rows[i].cells[1].text = context.get('engineer_name', '')
                table.rows[i].cells[3].text = context.get('calc_date', '')

        doc.add_paragraph()  # Spacing

    def _add_disclaimer(self, doc: Document, context: Dict[str, Any]):
        """Add legal disclaimer block"""
        doc.add_heading('DISCLAIMER', level=2)

        disclaimer_text = (
            "This calculation has been prepared by Beaver Bridges Ltd using BeaverCalc Studio "
            f"v{context.get('version', '1.0.0')}. The results are based solely on the input parameters "
            "provided and the assumptions stated herein. This document does not constitute a "
            "complete structural design and must be read in conjunction with the relevant "
            "drawings, specifications, and project documentation.\n\n"
            "The engineer of record is responsible for verifying that all inputs are correct, "
            "that the calculation method is appropriate for the intended application, and that "
            "the results are consistent with the overall design intent. Independent checking in "
            "accordance with the project's quality management procedures is mandatory before "
            "any results are relied upon for construction.\n\n"
            "Beaver Bridges Ltd accepts no liability for any loss, damage, or injury arising "
            "from the use of this calculation where the inputs have been incorrectly specified, "
            "where the calculation has not been independently checked, or where it has been "
            "applied outside its stated scope. All designs must comply with current Building "
            "Regulations and the relevant Eurocodes / British Standards."
        )

        p = doc.add_paragraph(disclaimer_text)
        for run in p.runs:
            run.font.size = Pt(8)
            run.font.italic = True

        doc.add_paragraph()  # Spacing

    def _add_appendices(self, doc: Document, context: Dict[str, Any]):
        """Add appendices section (references, code clauses)"""
        doc.add_heading('APPENDICES', level=2)

        # Appendix A — References
        doc.add_heading('Appendix A — References', level=2)
        design_code = context.get('design_code', 'EN 1993-1-1')
        references = [
            f'{design_code} — Eurocode design standard (as referenced)',
            'BS EN 1990 — Basis of structural design',
            'UK National Annex to the relevant Eurocode',
            'BeaverCalc Studio Technical Manual',
        ]
        for ref in references:
            p = doc.add_paragraph(ref, style='List Bullet')
            for run in p.runs:
                run.font.size = Pt(10)

        # Appendix B — Calculation log
        doc.add_heading('Appendix B — Calculation Log', level=2)

        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'
        log_items = [
            ('Run ID', str(context.get('run_id', 'N/A'))),
            ('Calculator', context.get('calculator_name', 'N/A')),
            ('Generated', context.get('calc_date', 'N/A')),
            ('Software Version', f"BeaverCalc Studio v{context.get('version', '1.0.0')}"),
        ]
        for i, (label, value) in enumerate(log_items):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
            table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

        doc.add_paragraph()  # Spacing

    def _add_footer(self, doc: Document, context: Dict[str, Any]):
        """Add report footer"""
        doc.add_page_break()
        
        footer_para = doc.add_paragraph()
        footer_para.add_run('BEAVER BRIDGES LTD\n').font.bold = True
        footer_para.add_run('Structural Engineering Excellence\n')
        footer_para.add_run('Company Registration: 09733378\n')
        footer_para.add_run('www.beaverbridges.co.uk\n\n')
        
        footer_para.add_run(f'This calculation has been prepared using BeaverCalc Studio v{context["version"]}\n')
        footer_para.add_run(f'© {context["current_year"]} Beaver Bridges Ltd. All rights reserved.\n\n')
        
        warning = footer_para.add_run('IMPORTANT: This calculation is valid only when bearing the original signature and company stamp.')
        warning.font.color.rgb = RGBColor(220, 53, 69)
        warning.font.bold = True
        
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer_para.runs:
            run.font.size = Pt(8)


# Convenience functions
def generate_pdf_report(run_data: Dict[str, Any], output_path: Optional[str] = None) -> bytes:
    """Generate PDF report"""
    generator = ReportGenerator()
    return generator.generate_pdf(run_data, output_path)


def generate_docx_report(run_data: Dict[str, Any], output_path: Optional[str] = None) -> bytes:
    """Generate DOCX report"""
    generator = ReportGenerator()
    return generator.generate_docx(run_data, output_path)
